import streamlit as st
import os
from docx import Document
from io import BytesIO
from input_forms.form_processing import insert_contract, insert_contract_details

# Directory to store generated contracts
CONTRACTS_DIR = "contracts_files"
if not os.path.exists(CONTRACTS_DIR):
    os.makedirs(CONTRACTS_DIR)

def replace_placeholders(doc, session_data):
    for para in doc.paragraphs:
        for key, value in session_data.items():
            placeholder = f"[{key}]"
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, str(value) if value else "")
    return doc

def run_contract_generation(contract_type):
    st.title("Generate Contract")

    if 'contract_data' not in st.session_state:
        st.error("No contract data available. Please fill in all contract details.")
        return

    # Normalize contract_type to match the database constraints
    contract_type = contract_type.lower().strip()
    normalized_contract_type = 'rental' if 'rental' in contract_type else 'sale' if 'sale' in contract_type else None

    if normalized_contract_type not in ['rental', 'sale']:
        st.error("Invalid contract type. Must be 'rental' or 'sale'.")
        return

    # Load template based on contract type
    template_path = "data/templates/updated_contract_rent.docx" if normalized_contract_type == "rental" else "data/templates/contract_sale.docx"

    if os.path.exists(template_path):
        doc = Document(template_path)
        contract_data = st.session_state['contract_data']

        # Replace placeholders in the document
        doc = replace_placeholders(doc, contract_data)

        # Show contract preview
        st.subheader("Contract Preview")
        preview_text = "\n".join([para.text for para in doc.paragraphs])
        st.text_area("Preview", value=preview_text, height=400)

        # Button to download contract
        if st.button("Download Contract"):
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            st.download_button(
                label="Download Your Contract",
                data=buffer,
                file_name=f"generated_{normalized_contract_type}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        # Button to save contract details to the database and save contract locally
        if st.button("Save Contract to Database"):
            # Determine the IDs for saving
            user_id_key = 'landlord_id' if normalized_contract_type == "rental" else 'seller_id'
            other_party_id_key = 'tenant_id' if normalized_contract_type == "rental" else 'buyer_id'

            # Insert contract into the database
            contract_id = insert_contract({
                'contract_type': normalized_contract_type,
                user_id_key: st.session_state.get(f'{user_id_key[:-3]}_cnp', 0),
                other_party_id_key: st.session_state.get(f'{other_party_id_key[:-3]}_cnp', 0)
            })
            insert_contract_details(contract_id, contract_data)

            # Save contract locally
            local_path = os.path.join(CONTRACTS_DIR, f"contract_{contract_id}.docx")
            doc.save(local_path)
            st.success(f"Contract saved successfully! File path: {local_path}")
    else:
        st.error("Template file not found.")
